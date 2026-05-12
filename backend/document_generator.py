import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_JUSTIFY

DARK   = RGBColor(0x1E, 0x29, 0x3B)
ACCENT = RGBColor(0x4F, 0x46, 0xE5)
BLACK  = RGBColor(0x0F, 0x17, 0x2A)
GRAY   = RGBColor(0x47, 0x55, 0x69)
PDF_DARK   = colors.HexColor("#1e293b")
PDF_ACCENT = colors.HexColor("#4f46e5")
PDF_BLACK  = colors.HexColor("#0f172a")
PDF_GRAY   = colors.HexColor("#475569")

def parse_lines(text):
    result = []
    for line in text.split("\n"):
        s = line.rstrip()
        if s.startswith("# "):       result.append({"type":"h1","text":s[2:].strip()})
        elif s.startswith("## "):    result.append({"type":"h2","text":s[3:].strip()})
        elif s.startswith("### "):   result.append({"type":"h3","text":s[4:].strip()})
        elif s.startswith(("- ","* ")): result.append({"type":"bullet","text":s[2:].strip()})
        elif s == "":                result.append({"type":"blank"})
        else:                        result.append({"type":"body","text":s})
    return result

def bold_runs(text, para, size, color, font="Calibri"):
    for p in re.split(r"(\*\*.*?\*\*)", text):
        r = para.add_run(p[2:-2] if p.startswith("**") and p.endswith("**") else p)
        r.bold = p.startswith("**")
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = font

def set_bottom_border(para, color_hex="4F46E5", sz=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def parse_cover_letter(text):
    data = {
        "sender_name":"","sender_address":"","sender_city":"",
        "sender_phone":"","sender_email":"",
        "date": datetime.now().strftime("%B %d, %Y"),
        "recipient_name":"","recipient_title":"","company_name":"","company_address":"",
        "salutation":"","paragraphs":[],"closing":"Sincerely,",
    }
    lines = [l.strip() for l in text.strip().split("\n")]
    non_empty = [l for l in lines if l]
    idx = 0

    def not_sal(l): return not l.lower().startswith("dear")

    if idx < len(non_empty) and not_sal(non_empty[idx]): data["sender_name"]    = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]): data["sender_address"] = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]) and not "@" in non_empty[idx] and not non_empty[idx][0].isdigit(): data["sender_city"] = non_empty[idx]; idx+=1
    if idx < len(non_empty) and non_empty[idx][0].isdigit(): data["sender_phone"] = non_empty[idx]; idx+=1
    if idx < len(non_empty) and "@" in non_empty[idx]: data["sender_email"] = non_empty[idx]; idx+=1
    if idx < len(non_empty) and any(m in non_empty[idx] for m in ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec","20"]): data["date"] = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]): data["recipient_name"]    = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]): data["recipient_title"]   = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]): data["company_name"]      = non_empty[idx]; idx+=1
    if idx < len(non_empty) and not_sal(non_empty[idx]): data["company_address"]   = non_empty[idx]; idx+=1

    sal_idx = next((i for i,l in enumerate(lines) if l.lower().startswith("dear")), None)
    if sal_idx is not None:
        data["salutation"] = lines[sal_idx]
        current = []
        for line in lines[sal_idx+1:]:
            low = line.lower().strip()
            if low in ["sincerely,","sincerely","regards,","regards","best regards,","warm regards,"]:
                if current: data["paragraphs"].append(" ".join(current)); current=[]
                data["closing"] = line; break
            elif line.strip() == "":
                if current: data["paragraphs"].append(" ".join(current)); current=[]
            else:
                current.append(line.strip())
        if current: data["paragraphs"].append(" ".join(current))
        data["paragraphs"] = [p for p in data["paragraphs"] if p.strip()]

    if not data["paragraphs"]:
        data["paragraphs"] = [text]
    return data

def generate_cover_letter_docx(text, output_path, company):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.2)
        sec.left_margin = sec.right_margin = Inches(1.25)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    d = parse_cover_letter(text)

    def L(txt, bold=False, size=12, space_after=0, italic=False, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(15)
        r = p.add_run(txt)
        r.bold=bold; r.italic=italic
        r.font.name="Times New Roman"; r.font.size=Pt(size)
        r.font.color.rgb = color if color else BLACK
        return p

    def B(n=10): p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(n)

    if d["sender_name"]:    L(d["sender_name"], bold=True)
    if d["sender_address"]: L(d["sender_address"])
    if d["sender_city"]:    L(d["sender_city"])
    if d["sender_phone"]:   L(d["sender_phone"])
    if d["sender_email"]:   L(d["sender_email"])
    B(14)
    L(d["date"])
    B(14)
    if d["recipient_name"]:  L(d["recipient_name"], bold=True)
    if d["recipient_title"]: L(d["recipient_title"])
    if d["company_name"]:    L(d["company_name"])
    if d["company_address"]: L(d["company_address"])
    B(14)
    L(d["salutation"] or "Dear Hiring Manager,")
    B(10)

    for pt in d["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(10)
        p.paragraph_format.line_spacing = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(pt.strip())
        r.font.name="Times New Roman"; r.font.size=Pt(12); r.font.color.rgb=BLACK

    B(14)
    L(d["closing"] or "Sincerely,")
    B(36)
    if d["sender_name"]:
        L(d["sender_name"])
        p=doc.add_paragraph()
        r=p.add_run(d["sender_name"]); r.italic=True
        r.font.name="Times New Roman"; r.font.size=Pt(12); r.font.color.rgb=GRAY
    B(8)
    L("Enclosures")
    doc.save(output_path)

def generate_cover_letter_pdf(text, output_path, company):
    doc = SimpleDocTemplate(output_path, pagesize=letter,
        rightMargin=inch*1.25, leftMargin=inch*1.25,
        topMargin=inch*1.2, bottomMargin=inch*1.2)
    N  = ParagraphStyle("N",  fontName="Times-Roman",  fontSize=12, leading=16, textColor=PDF_BLACK, spaceAfter=0)
    B  = ParagraphStyle("B",  fontName="Times-Bold",   fontSize=12, leading=16, textColor=PDF_BLACK, spaceAfter=0)
    BD = ParagraphStyle("BD", fontName="Times-Roman",  fontSize=12, leading=18, textColor=PDF_BLACK, spaceAfter=10, alignment=TA_JUSTIFY)
    IT = ParagraphStyle("IT", fontName="Times-Italic", fontSize=12, leading=16, textColor=PDF_GRAY,  spaceAfter=0)
    d  = parse_cover_letter(text)
    s  = []
    sf = lambda t: t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    SP = lambda n=8: Spacer(1,n)

    if d["sender_name"]:    s+=[Paragraph(sf(d["sender_name"]),B)]
    if d["sender_address"]: s+=[Paragraph(sf(d["sender_address"]),N)]
    if d["sender_city"]:    s+=[Paragraph(sf(d["sender_city"]),N)]
    if d["sender_phone"]:   s+=[Paragraph(sf(d["phone"] if "phone" in d else d["sender_phone"]),N)]
    if d["sender_email"]:   s+=[Paragraph(sf(d["sender_email"]),N)]
    s+=[SP(14),Paragraph(sf(d["date"]),N),SP(14)]
    if d["recipient_name"]:  s+=[Paragraph(sf(d["recipient_name"]),B)]
    if d["recipient_title"]: s+=[Paragraph(sf(d["recipient_title"]),N)]
    if d["company_name"]:    s+=[Paragraph(sf(d["company_name"]),N)]
    if d["company_address"]: s+=[Paragraph(sf(d["company_address"]),N)]
    s+=[SP(14),Paragraph(sf(d["salutation"] or "Dear Hiring Manager,"),N),SP(10)]
    for pt in d["paragraphs"]:
        if pt.strip(): s.append(Paragraph(sf(pt.strip()),BD))
    s+=[SP(14),Paragraph(sf(d["closing"] or "Sincerely,"),N),SP(40)]
    if d["sender_name"]:
        s+=[Paragraph(sf(d["sender_name"]),N),Paragraph(f"<i>{sf(d['sender_name'])}</i>",IT)]
    s+=[SP(10),Paragraph("Enclosures",N)]
    doc.build(s)

def generate_resume_docx(text, output_path, company):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=sec.bottom_margin=Inches(0.9)
        sec.left_margin=sec.right_margin=Inches(1.0)
    doc.styles["Normal"].font.name="Calibri"
    doc.styles["Normal"].font.size=Pt(10.5)
    for item in parse_lines(text):
        t,txt=item["type"],item.get("text","")
        if t=="blank":
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        elif t=="h1":
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
            r=p.add_run(txt.upper()); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=DARK; r.font.name="Calibri"
            set_bottom_border(p)
        elif t=="h2":
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
            r=p.add_run(txt); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=ACCENT; r.font.name="Calibri"
        elif t=="h3":
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
            r=p.add_run(txt); r.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=DARK; r.font.name="Calibri"
        elif t=="bullet":
            p=doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after=Pt(2); p.paragraph_format.left_indent=Inches(0.2)
            bold_runs(txt,p,10.5,BLACK)
        elif t=="body":
            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
            bold_runs(txt,p,10.5,BLACK)
    doc.save(output_path)

def generate_resume_pdf(text, output_path, company):
    doc=SimpleDocTemplate(output_path,pagesize=letter,
        rightMargin=inch,leftMargin=inch,topMargin=inch*0.9,bottomMargin=inch*0.9)
    h1=ParagraphStyle("h1",fontName="Helvetica-Bold",fontSize=14,textColor=PDF_DARK,spaceBefore=12,spaceAfter=4)
    h2=ParagraphStyle("h2",fontName="Helvetica-Bold",fontSize=11,textColor=PDF_ACCENT,spaceBefore=8,spaceAfter=3)
    h3=ParagraphStyle("h3",fontName="Helvetica-Bold",fontSize=10,textColor=PDF_DARK,spaceBefore=5,spaceAfter=2)
    bd=ParagraphStyle("bd",fontName="Helvetica",fontSize=10,textColor=PDF_BLACK,leading=15,spaceAfter=3,alignment=TA_JUSTIFY)
    bl=ParagraphStyle("bl",fontName="Helvetica",fontSize=10,textColor=PDF_BLACK,leading=14,leftIndent=14,spaceAfter=2)
    story=[]
    for item in parse_lines(text):
        t,txt=item["type"],item.get("text","")
        safe=re.sub(r"\*\*(.*?)\*\*",r"<b>\1</b>",txt.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))
        if   t=="blank":  story.append(Spacer(1,4))
        elif t=="h1":     story+=[Paragraph(safe,h1),HRFlowable(width="100%",thickness=1.5,color=PDF_ACCENT,spaceAfter=4)]
        elif t=="h2":     story.append(Paragraph(safe,h2))
        elif t=="h3":     story.append(Paragraph(safe,h3))
        elif t=="bullet": story.append(Paragraph("• "+safe,bl))
        elif t=="body":   story.append(Paragraph(safe,bd))
    doc.build(story)

def generate_email_docx(text, output_path, company):
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=sec.bottom_margin=Inches(1.0)
        sec.left_margin=sec.right_margin=Inches(1.25)
    doc.styles["Normal"].font.name="Calibri"
    doc.styles["Normal"].font.size=Pt(11)
    for line in text.strip().split("\n"):
        s=line.strip()
        p=doc.add_paragraph()
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.line_spacing=Pt(16)
        r=p.add_run(s)
        if s.lower().startswith("subject:"): r.bold=True; r.font.size=Pt(12); r.font.color.rgb=ACCENT
        else: r.font.size=Pt(11); r.font.color.rgb=BLACK
        r.font.name="Calibri"
    doc.save(output_path)

def generate_email_pdf(text, output_path, company):
    doc=SimpleDocTemplate(output_path,pagesize=letter,
        rightMargin=inch*1.25,leftMargin=inch*1.25,topMargin=inch,bottomMargin=inch)
    subj=ParagraphStyle("S",fontName="Helvetica-Bold",fontSize=12,textColor=PDF_ACCENT,spaceAfter=10)
    body=ParagraphStyle("B",fontName="Helvetica",fontSize=11,textColor=PDF_BLACK,leading=17,spaceAfter=8,alignment=TA_JUSTIFY)
    story=[]
    for line in text.strip().split("\n"):
        s=line.strip()
        safe=s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        if   s.lower().startswith("subject:"): story.append(Paragraph(safe,subj))
        elif s=="":                             story.append(Spacer(1,6))
        else:                                   story.append(Paragraph(safe,body))
    doc.build(story)

def generate_docx(content, output_path, doc_type, company):
    if doc_type=="coverLetter": generate_cover_letter_docx(content,output_path,company)
    elif doc_type=="email":     generate_email_docx(content,output_path,company)
    else:                       generate_resume_docx(content,output_path,company)

def generate_pdf(content, output_path, doc_type, company):
    if doc_type=="coverLetter": generate_cover_letter_pdf(content,output_path,company)
    elif doc_type=="email":     generate_email_pdf(content,output_path,company)
    else:                       generate_resume_pdf(content,output_path,company)


